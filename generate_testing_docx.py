#!/usr/bin/env python3
"""
Generate DOCX (Microsoft Word) document for AI Agent OTEL Testing Manual
"""

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("ERROR: python-docx not installed")
    print("Please install it: pip install python-docx")
    exit(1)

import os
from datetime import datetime

def add_heading(doc, text, level=1):
    """Add a heading with formatting"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, style=None):
    """Add a paragraph with optional style"""
    para = doc.add_paragraph(text)
    if style:
        para.style = style
    return para

def add_code_block(doc, code):
    """Add a code block with monospace font"""
    para = doc.add_paragraph(code)
    para.style = 'Normal'
    for run in para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 0, 128)
    para.paragraph_format.left_indent = Inches(0.5)
    return para

def add_table_from_data(doc, headers, rows):
    """Add a formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Add headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Add data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)

    return table

def create_testing_manual():
    """Create the complete testing manual DOCX"""

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title Page
    title = doc.add_heading('AI Agent OTEL Integration', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('Testing & Validation Manual', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Document info
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.add_run(f'Document Version: 1.0\n').bold = True
    info_para.add_run(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n')
    info_para.add_run(f'Project: agentic-ai-integration-490716\n')

    doc.add_page_break()

    # Table of Contents (manual)
    add_heading(doc, 'Table of Contents', level=1)
    toc_items = [
        '1. Quick Start',
        '2. Overview',
        '3. Prerequisites',
        '4. Test Scripts Reference',
        '5. Step-by-Step Testing Guide',
        '6. Portal26 Verification',
        '7. Troubleshooting',
        '8. Success Criteria',
        '9. Appendix: Service Information',
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # Section 1: Quick Start
    add_heading(doc, '1. Quick Start', level=1)

    add_paragraph(doc, 'Three ways to validate your OTEL integration:')
    doc.add_paragraph()

    add_heading(doc, 'Option A: Automated Testing (2 minutes)', level=2)
    add_paragraph(doc, 'Run the automated validation suite:')
    add_code_block(doc, 'python run_validation.py')
    add_paragraph(doc, 'This will test all components and generate a comprehensive report.')
    doc.add_paragraph()

    add_heading(doc, 'Option B: Manual Testing (5 minutes)', level=2)
    add_paragraph(doc, 'Step 1: Test OTLP endpoints')
    add_code_block(doc, 'python test_otel_send.py')
    add_paragraph(doc, 'Step 2: Send test request to AI agent')
    add_code_block(doc, 'TEST_ID="TEST-$(date +%s)"\nTOKEN=$(gcloud auth print-identity-token)\n\ncurl -X POST https://ai-agent-961756870884.us-central1.run.app/chat \\\n  -H "Authorization: Bearer $TOKEN" \\\n  -H "Content-Type: application/json" \\\n  -d "{\\"message\\": \\"$TEST_ID: What is the weather in Tokyo?\\"}"')
    add_paragraph(doc, 'Step 3: Search for your TEST_ID in Portal26 dashboard')
    doc.add_paragraph()

    doc.add_page_break()

    # Section 2: Overview
    add_heading(doc, '2. Overview', level=1)

    add_paragraph(doc, 'This manual provides complete instructions for testing and validating the AI Agent OpenTelemetry (OTEL) integration with Portal26.')
    doc.add_paragraph()

    add_heading(doc, 'What Gets Tested', level=2)
    test_items = [
        'Cloud Run service deployment and health',
        'OTEL configuration (traces, metrics, logs)',
        'Portal26 OTLP endpoint connectivity',
        'Data export from AI Agent to Portal26',
        'End-to-end telemetry flow',
    ]
    for item in test_items:
        doc.add_paragraph(f'✓ {item}', style='List Bullet')

    doc.add_paragraph()

    add_heading(doc, 'Architecture Overview', level=2)
    add_paragraph(doc, 'Data Flow:')
    flow = '''User Request
    ↓
AI Agent (Cloud Run)
    ↓
OpenTelemetry SDK
    ↓
OTLP Exporters
    ↓
Portal26 OTLP Endpoints
    ↓
Portal26 Dashboard'''
    add_code_block(doc, flow)

    doc.add_page_break()

    # Section 3: Prerequisites
    add_heading(doc, '3. Prerequisites', level=1)

    add_heading(doc, 'Required Tools', level=2)
    prereq_table = [
        ['Tool', 'Purpose', 'Installation'],
        ['gcloud CLI', 'Google Cloud authentication', 'cloud.google.com/sdk/docs/install'],
        ['Python 3.7+', 'Run test scripts', 'python.org/downloads'],
        ['requests', 'HTTP library', 'pip install requests'],
        ['opentelemetry-sdk', 'OTEL SDK', 'pip install opentelemetry-sdk'],
    ]
    add_table_from_data(doc, prereq_table[0], prereq_table[1:])

    doc.add_paragraph()

    add_heading(doc, 'Authentication Setup', level=2)
    add_paragraph(doc, 'Before testing, authenticate with Google Cloud:')
    add_code_block(doc, 'gcloud auth login\ngcloud auth application-default login\ngcloud config set project agentic-ai-integration-490716')

    doc.add_page_break()

    # Section 4: Test Scripts
    add_heading(doc, '4. Test Scripts Reference', level=1)

    scripts_table = [
        ['Script', 'Purpose', 'Runtime', 'Output'],
        ['run_validation.py', 'Complete automated validation', '1-2 min', 'Full report + JSON file'],
        ['test_otel_send.py', 'Test OTLP endpoints directly', '10 sec', 'HTTP 200 responses'],
        ['verify_otel_export.py', 'End-to-end verification', '30 sec', 'Verification report'],
        ['capture_otel_export.py', 'Real-time export proof', '15 sec', 'Test ID for Portal26'],
    ]
    add_table_from_data(doc, scripts_table[0], scripts_table[1:])

    doc.add_paragraph()

    add_heading(doc, 'Script Descriptions', level=2)

    add_heading(doc, 'run_validation.py (Recommended)', level=3)
    add_paragraph(doc, 'Comprehensive automated test suite that:')
    validation_steps = [
        'Checks gcloud authentication',
        'Verifies Cloud Run deployment',
        'Tests OTEL configuration',
        'Sends test requests to AI agent',
        'Tests all OTLP endpoints',
        'Checks logs for errors',
        'Generates validation report',
    ]
    for step in validation_steps:
        doc.add_paragraph(f'• {step}', style='List Bullet')

    doc.add_paragraph()
    add_code_block(doc, 'python run_validation.py')

    doc.add_page_break()

    # Section 5: Step-by-Step Testing
    add_heading(doc, '5. Step-by-Step Testing Guide', level=1)

    add_heading(doc, 'Phase 1: Verify Deployment', level=2)

    add_heading(doc, 'Step 1.1: Check Service Status', level=3)
    add_code_block(doc, 'gcloud run services describe ai-agent \\\n  --region=us-central1 \\\n  --format="value(status.url)"')

    add_paragraph(doc, 'Expected Output:')
    add_code_block(doc, 'https://ai-agent-961756870884.us-central1.run.app')

    doc.add_paragraph()

    add_heading(doc, 'Step 1.2: Test Status Endpoint', level=3)
    add_code_block(doc, 'TOKEN=$(gcloud auth print-identity-token)\n\ncurl -H "Authorization: Bearer $TOKEN" \\\n  https://ai-agent-961756870884.us-central1.run.app/status')

    add_paragraph(doc, 'Expected Output:')
    add_code_block(doc, '{\n  "agent_mode": "manual",\n  "manual_agent_enabled": true,\n  "adk_agent_enabled": false\n}')

    doc.add_paragraph()

    add_heading(doc, 'Phase 2: Test OTLP Endpoints', level=2)

    add_paragraph(doc, 'Run the endpoint test script:')
    add_code_block(doc, 'python test_otel_send.py')

    add_paragraph(doc, 'Expected Output:')
    expected_output = '''[OK] Traces endpoint accepting data!
Status: 200
Response: {"partialSuccess":{}}

[OK] Metrics endpoint accepting data!
Status: 200
Response: {"partialSuccess":{}}

[OK] Logs endpoint accepting data!
Status: 200
Response: {"partialSuccess":{}}'''
    add_code_block(doc, expected_output)

    doc.add_paragraph()

    add_heading(doc, 'Phase 3: Send Test Request', level=2)

    add_paragraph(doc, 'Generate a unique test ID and send a request:')
    test_request = '''TEST_ID="TEST-$(date +%s)"
TOKEN=$(gcloud auth print-identity-token)

curl -X POST https://ai-agent-961756870884.us-central1.run.app/chat \\
  -H "Authorization: Bearer $TOKEN" \\
  -H "Content-Type: application/json" \\
  -d "{\\"message\\": \\"$TEST_ID: What is the weather in Tokyo?\\"}"

echo "Test ID: $TEST_ID"'''
    add_code_block(doc, test_request)

    add_paragraph(doc, 'Save the Test ID for Portal26 verification!')

    doc.add_page_break()

    # Section 6: Portal26 Verification
    add_heading(doc, '6. Portal26 Verification', level=1)

    add_heading(doc, 'Step 6.1: Access Dashboard', level=2)
    portal_info = [
        ['Field', 'Value'],
        ['URL', 'https://portal26.in'],
        ['Username', 'titaniam'],
        ['Password', 'helloworld'],
        ['Tenant', 'tenant1'],
        ['User ID', 'relusys'],
    ]
    add_table_from_data(doc, portal_info[0], portal_info[1:])

    doc.add_paragraph()

    add_heading(doc, 'Step 6.2: Apply Filters', level=2)
    add_paragraph(doc, 'In the Portal26 dashboard, apply these filters:')
    filter_items = [
        'Service Name: ai-agent',
        'Tenant ID: tenant1',
        'User ID: relusys',
        'Environment: production',
        'Time Range: Last 15 minutes',
    ]
    for item in filter_items:
        doc.add_paragraph(f'• {item}', style='List Bullet')

    doc.add_paragraph()

    add_heading(doc, 'Step 6.3: Search for Test Traces', level=2)
    add_paragraph(doc, 'Use the test ID from Step 3 to search for your specific trace.')
    add_paragraph(doc, 'You should see:')
    trace_details = [
        'Trace with span name: "agent_chat"',
        'Duration: ~1-2 seconds',
        'Status: Success',
        'Attributes: user.message, agent.success, agent_mode',
    ]
    for detail in trace_details:
        doc.add_paragraph(f'✓ {detail}', style='List Bullet')

    doc.add_paragraph()

    add_heading(doc, 'Step 6.4: Verify Metrics', level=2)
    metrics_table = [
        ['Metric Name', 'Type', 'Expected Values'],
        ['agent_requests_total', 'Counter', 'Increasing with each request'],
        ['agent_response_time_seconds', 'Histogram', 'Distribution of response times (1-3s)'],
    ]
    add_table_from_data(doc, metrics_table[0], metrics_table[1:])

    doc.add_paragraph()

    add_heading(doc, 'Step 6.5: Check Logs', level=2)
    add_paragraph(doc, 'Look for log entries such as:')
    log_examples = [
        'Chat request received: What is the weather...',
        'Agent processing',
        'Chat request completed successfully in 1.23s',
    ]
    for log in log_examples:
        doc.add_paragraph(f'• {log}', style='List Bullet')

    doc.add_page_break()

    # Section 7: Troubleshooting
    add_heading(doc, '7. Troubleshooting', level=1)

    add_heading(doc, 'Issue 1: Authentication Errors', level=2)
    add_paragraph(doc, 'Symptom: 403 Forbidden or authentication failures')
    add_paragraph(doc, 'Solution:')
    add_code_block(doc, 'gcloud auth login\ngcloud auth application-default login')

    doc.add_paragraph()

    add_heading(doc, 'Issue 2: OTEL Exporters Not Configured', level=2)
    add_paragraph(doc, 'Symptom: No telemetry logs in Cloud Run')
    add_paragraph(doc, 'Solution: Verify environment variables')
    add_code_block(doc, 'gcloud run services describe ai-agent \\\n  --region=us-central1 \\\n  --format="yaml(spec.template.spec.containers[0].env)" | grep OTEL')

    doc.add_paragraph()

    add_heading(doc, 'Issue 3: No Data in Portal26', level=2)
    add_paragraph(doc, 'Possible causes and solutions:')
    portal26_issues = [
        'Time Range: Expand to "Last 1 hour"',
        'Filters: Remove all filters and search again',
        'Tenant/User: Verify correct tenant (tenant1)',
        'Delay: Wait 2-3 minutes for data propagation',
        'Dashboard: Check if there\'s a different view for OTLP data',
    ]
    for issue in portal26_issues:
        doc.add_paragraph(f'• {issue}', style='List Bullet')

    doc.add_paragraph()

    add_heading(doc, 'Issue 4: Portal26 Endpoint Errors', level=2)
    add_paragraph(doc, 'Symptom: 401, 403, or 500 errors')
    add_paragraph(doc, 'Solution: Verify authentication')
    add_code_block(doc, 'echo "dGl0YW5pYW06aGVsbG93b3JsZA==" | base64 -d')
    add_paragraph(doc, 'Should output: titaniam:helloworld')

    doc.add_page_break()

    # Section 8: Success Criteria
    add_heading(doc, '8. Success Criteria', level=1)

    add_paragraph(doc, 'Your OTEL integration is successful when all these criteria are met:')
    doc.add_paragraph()

    success_criteria = [
        ('Deployment', [
            'Cloud Run service is running',
            'Service URL is accessible',
            'Status endpoint returns JSON',
            'Chat endpoint processes requests',
        ]),
        ('Configuration', [
            'OTEL_TRACES_EXPORTER=otlp',
            'OTEL_METRICS_EXPORTER=otlp',
            'OTEL_LOGS_EXPORTER=otlp',
            'Portal26 endpoint configured',
            'Authentication configured',
        ]),
        ('Export', [
            'Traces endpoint returns 200 OK',
            'Metrics endpoint returns 200 OK',
            'Logs endpoint returns 200 OK',
            'All exporters initialized in logs',
            'No export errors found',
        ]),
        ('Portal26', [
            'Dashboard accessible',
            'Can filter by service: ai-agent',
            'Test traces visible within 2-3 minutes',
            'Metrics updating every 1 second',
            'Logs streaming every 500ms',
            'Can search by test ID',
        ]),
    ]

    for category, items in success_criteria:
        add_heading(doc, category, level=2)
        for item in items:
            doc.add_paragraph(f'✓ {item}', style='List Bullet')
        doc.add_paragraph()

    doc.add_page_break()

    # Section 9: Appendix
    add_heading(doc, '9. Appendix: Service Information', level=1)

    add_heading(doc, 'Service Details', level=2)
    service_info = [
        ['Property', 'Value'],
        ['Service Name', 'ai-agent'],
        ['Project ID', 'agentic-ai-integration-490716'],
        ['Region', 'us-central1'],
        ['Service URL', 'https://ai-agent-961756870884.us-central1.run.app'],
    ]
    add_table_from_data(doc, service_info[0], service_info[1:])

    doc.add_paragraph()

    add_heading(doc, 'Portal26 Configuration', level=2)
    portal26_config = [
        ['Property', 'Value'],
        ['OTLP Endpoint', 'https://otel-tenant1.portal26.in:4318'],
        ['Dashboard URL', 'https://portal26.in'],
        ['Tenant ID', 'tenant1'],
        ['User ID', 'relusys'],
        ['Username', 'titaniam'],
        ['Password', 'helloworld'],
    ]
    add_table_from_data(doc, portal26_config[0], portal26_config[1:])

    doc.add_paragraph()

    add_heading(doc, 'OTLP Endpoints', level=2)
    endpoints_info = [
        ['Endpoint', 'URL', 'Purpose'],
        ['Traces', 'https://otel-tenant1.portal26.in:4318/v1/traces', 'Receive trace data'],
        ['Metrics', 'https://otel-tenant1.portal26.in:4318/v1/metrics', 'Receive metrics data'],
        ['Logs', 'https://otel-tenant1.portal26.in:4318/v1/logs', 'Receive log data'],
    ]
    add_table_from_data(doc, endpoints_info[0], endpoints_info[1:])

    doc.add_paragraph()

    add_heading(doc, 'Telemetry Configuration', level=2)
    telemetry_config = [
        ['Property', 'Value'],
        ['Traces Export', 'OTLP HTTP'],
        ['Metrics Export', 'OTLP HTTP'],
        ['Logs Export', 'OTLP HTTP'],
        ['Protocol', 'http/protobuf'],
        ['Metrics Interval', '1000ms (1 second)'],
        ['Logs Interval', '500ms (0.5 seconds)'],
    ]
    add_table_from_data(doc, telemetry_config[0], telemetry_config[1:])

    doc.add_paragraph()

    add_heading(doc, 'Resource Attributes', level=2)
    add_paragraph(doc, 'Every telemetry signal includes these resource attributes:')
    resource_attrs = [
        ['Attribute', 'Value', 'Purpose'],
        ['service.name', 'ai-agent', 'Service identifier'],
        ['portal26.user.id', 'relusys', 'Portal26 user'],
        ['portal26.tenant_id', 'tenant1', 'Portal26 tenant'],
        ['service.version', '1.0', 'Service version'],
        ['deployment.environment', 'production', 'Environment'],
    ]
    add_table_from_data(doc, resource_attrs[0], resource_attrs[1:])

    doc.add_paragraph()

    add_heading(doc, 'Google Cloud Console Links', level=2)
    console_links = [
        ['Resource', 'URL'],
        ['Cloud Run', 'https://console.cloud.google.com/run/detail/us-central1/ai-agent'],
        ['Logs', 'https://console.cloud.google.com/logs/query?project=agentic-ai-integration-490716'],
        ['Monitoring', 'https://console.cloud.google.com/monitoring?project=agentic-ai-integration-490716'],
    ]
    add_table_from_data(doc, console_links[0], console_links[1:])

    doc.add_paragraph()

    # Footer
    doc.add_page_break()
    add_heading(doc, 'Document Information', level=1)
    doc_info = [
        ['Property', 'Value'],
        ['Document Version', '1.0'],
        ['Generated', datetime.now().strftime('%Y-%m-%d %H:%M:%S')],
        ['Project', 'agentic-ai-integration-490716'],
        ['Status', '✓ Validated & Tested'],
    ]
    add_table_from_data(doc, doc_info[0], doc_info[1:])

    doc.add_paragraph()
    add_paragraph(doc, 'For the latest updates and additional documentation, see the project repository.')

    return doc

# Main execution
print("=" * 70)
print("AI Agent OTEL Testing Manual - DOCX Generator")
print("=" * 70)
print()

print("Generating DOCX document...")
try:
    doc = create_testing_manual()

    output_file = 'AI_Agent_OTEL_Testing_Manual.docx'
    doc.save(output_file)

    file_size = os.path.getsize(output_file)
    print(f"\n[SUCCESS] Document created: {output_file}")
    print(f"File size: {file_size:,} bytes ({file_size/1024:.1f} KB)")
    print()
    print("Document includes:")
    print("  ✓ Complete testing procedures")
    print("  ✓ Step-by-step instructions")
    print("  ✓ Troubleshooting guide")
    print("  ✓ Service information")
    print("  ✓ Portal26 configuration")
    print("  ✓ Code examples and commands")
    print("  ✓ Tables and formatted content")
    print()
    print(f"Open with Microsoft Word: {output_file}")
    print()

except Exception as e:
    print(f"\n[ERROR] Failed to create document: {e}")
    print()
    import traceback
    traceback.print_exc()
